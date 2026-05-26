# -*- coding: utf-8 -*-
"""
主要用途：
  将 Markdown 格式的产品需求文档（PRD）转换为 Word（.docx），便于评审与对外分发。

默认输入/输出：
  admin/docs/提拆派报价需求文档.md  →  admin/docs/提拆派报价需求文档.docx

用法：
  python admin/docs/_md_to_docx.py

说明：
  支持标题、表格、列表、加粗、图片引用 ![说明](images/模块/xx.png)。
  Markdown 的 **加粗** 在 docx 中转为 Word 粗体，不保留 ** 字符（含表格单元格）。
  第 3 章流程图须为 PNG（先运行 _render_mermaid.py）；界面截图（_capture_prd_screenshots.py）。
  附录中的 Mermaid 代码块仍以等宽文本保留，不会自动渲染为图。
  若 docx 被 Word 占用，会改写到「提拆派报价需求文档-最新.docx」。
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

# Word A4 默认页边距下，正文区约 6.2" 宽；流程图偏高时需限制高度
PAGE_CONTENT_WIDTH_IN = 5.25
PAGE_CONTENT_MAX_HEIGHT_IN = 6.5
UI_SCREENSHOT_WIDTH_IN = 5.25
FLOWCHART_MAX_WIDTH_IN = 5.0
FLOWCHART_MAX_HEIGHT_IN = 6.2
# 3.1 核心业务流程图（竖向）在 Word 中缩小，便于与 §3 标题同页展示
FLOWCHART_CORE_BUSINESS_MAX_WIDTH_IN = 2.55
FLOWCHART_CORE_BUSINESS_MAX_HEIGHT_IN = 3.1

FONT_NAME = "微软雅黑"


def md_plain(text: str) -> str:
    """去掉 Markdown 加粗标记，保留文字（用于标题等）。"""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _clear_paragraph(paragraph) -> None:
    p_el = paragraph._element
    for child in list(p_el):
        if child.tag.endswith("}r"):
            p_el.remove(child)


def apply_md_inline(
    paragraph,
    text: str,
    *,
    font_size_pt: float = 10,
    default_bold: bool = False,
    italic: bool = False,
) -> None:
    """将 **粗体** 转为 Word 加粗 run，不输出 ** 字符。"""
    _clear_paragraph(paragraph)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = default_bold
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            content = part[2:-2]
            bold = True
        else:
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", part)
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(font_size_pt)


def set_cell_md_text(cell, text: str, *, header_row: bool = False) -> None:
    cell.text = ""
    apply_md_inline(
        cell.paragraphs[0],
        (text or "").strip(),
        font_size_pt=10,
        default_bold=header_row,
    )


def set_cell_shading(cell, fill_hex: str):
    from docx.oxml import OxmlElement

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(col_count):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            set_cell_md_text(cell, text, header_row=(i == 0))
    tail = doc.add_paragraph()
    tail.paragraph_format.space_after = Pt(2)


def compact_heading_styles(doc):
    """收紧标题段前段后、取消标题前分页，减少第 3 章等大段空白。"""
    specs = (
        ("Heading 1", Pt(12), Pt(6), False),
        ("Heading 2", Pt(8), Pt(4), True),
        ("Heading 3", Pt(6), Pt(2), True),
        ("Heading 4", Pt(4), Pt(2), True),
    )
    for name, before, after, keep_next in specs:
        if name not in doc.styles:
            continue
        pf = doc.styles[name].paragraph_format
        pf.page_break_before = False
        pf.space_before = before
        pf.space_after = after
        pf.keep_with_next = keep_next


def _png_pixel_size(path: Path) -> tuple[int, int] | None:
    try:
        with path.open("rb") as f:
            if f.read(8) != b"\x89PNG\r\n\x1a\n":
                return None
            f.read(4)
            f.read(4)
            chunk = f.read(8)
            if chunk[:4] != b"IHDR":
                return None
            w, h = int.from_bytes(chunk[4:8], "big"), int.from_bytes(chunk[8:12], "big")
            return w, h
    except OSError:
        return None


def _is_flowchart_image(src: str, stem: str) -> bool:
    name = Path(src).name
    if re.match(r"0[5-9]-", name):
        return True
    return any(k in stem for k in ("流程", "流转", "flowchart", "mermaid"))


def _is_core_business_flowchart(src: str) -> bool:
    return "05-核心业务流程" in Path(src).name


def _fit_picture_emu(img_path: Path, src: str) -> tuple[int, int | None]:
    """按页面可显示区域计算插入 Word 的宽高（EMU），保持比例。"""
    px = _png_pixel_size(img_path)
    if not px or px[0] <= 0 or px[1] <= 0:
        max_w = FLOWCHART_MAX_WIDTH_IN if _is_flowchart_image(src, img_path.stem) else UI_SCREENSHOT_WIDTH_IN
        return int(Inches(max_w)), None

    w_px, h_px = px
    aspect = h_px / w_px
    if _is_core_business_flowchart(src):
        max_w_in, max_h_in = FLOWCHART_CORE_BUSINESS_MAX_WIDTH_IN, FLOWCHART_CORE_BUSINESS_MAX_HEIGHT_IN
    elif _is_flowchart_image(src, img_path.stem):
        max_w_in, max_h_in = FLOWCHART_MAX_WIDTH_IN, FLOWCHART_MAX_HEIGHT_IN
    else:
        max_w_in, max_h_in = UI_SCREENSHOT_WIDTH_IN, PAGE_CONTENT_MAX_HEIGHT_IN

    w_in = min(max_w_in, PAGE_CONTENT_WIDTH_IN)
    h_in = w_in * aspect
    if h_in > max_h_in:
        h_in = max_h_in
        w_in = h_in / aspect
    return int(Inches(w_in)), int(Inches(h_in))


def add_markdown_image(doc, md_path: Path, alt: str, src: str):
    img_path = Path(src)
    if not img_path.is_absolute():
        img_path = (md_path.parent / img_path).resolve()
    if not img_path.is_file():
        p = doc.add_paragraph()
        run = p.add_run(f"[图片缺失: {src}]")
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = None
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    width_emu, height_emu = _fit_picture_emu(img_path, src)
    run = p.add_run()
    if height_emu:
        run.add_picture(str(img_path), width=width_emu, height=height_emu)
    else:
        run.add_picture(str(img_path), width=width_emu)
    # 图题由 md 中 *图 x.x* 行提供，此处不再重复插入 alt 段落


def parse_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_doc_font(doc)
    compact_heading_styles(doc)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    i = 0
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and "|" in stripped[1:]:
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        else:
            flush_table()

        if stripped == "---":
            # 跳过分隔线，避免 Word 中多余空行或整页空白
            i += 1
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(md_plain(stripped[2:].strip()), level=1)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(md_plain(stripped[3:].strip()), level=2)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(md_plain(stripped[4:].strip()), level=3)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
        if stripped.startswith("###### "):
            h = doc.add_heading(md_plain(stripped[7:].strip()), level=5)
            h.paragraph_format.space_before = Pt(4)
            h.paragraph_format.space_after = Pt(2)
            i += 1
            continue
        if stripped.startswith("##### "):
            h = doc.add_heading(md_plain(stripped[6:].strip()), level=4)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(md_plain(stripped[5:].strip()), level=4)
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(0.5)
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            apply_md_inline(p, stripped[2:].strip(), font_size_pt=10.5, italic=True)
            i += 1
            continue

        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if img_m:
            add_markdown_image(doc, md_path, img_m.group(1), img_m.group(2))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            apply_md_inline(p, stripped[2:].strip(), font_size_pt=10.5)
            i += 1
            continue

        cap_m = re.match(r"^\*([^*]+)\*$", stripped)
        if cap_m:
            cap = doc.add_paragraph(cap_m.group(1).strip())
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(0)
            cap.paragraph_format.space_after = Pt(4)
            for run in cap.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(9)
                run.italic = True
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            # 使用正文段落 + 显式序号，避免 Word 全局连续编号导致 4.3 节从 9 开始
            p = doc.add_paragraph()
            apply_md_inline(p, f"{m.group(1)}. {m.group(2)}", font_size_pt=10.5)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        apply_md_inline(p, stripped, font_size_pt=10.5)
        i += 1

    flush_table()
    doc.save(docx_path)


if __name__ == "__main__":
    import shutil
    import sys

    base = Path(__file__).resolve().parent
    md = base / "提拆派报价需求文档.md"
    out = base / "提拆派报价需求文档.docx"
    tmp = base / "提拆派报价需求文档._gen.docx"

    parse_md_to_docx(md, tmp)
    try:
        shutil.copy2(tmp, out)
        tmp.unlink(missing_ok=True)
        print(f"已更新: {out}")
    except PermissionError:
        # 不另建「-最新」「-流程图优化」等副本，避免目录里堆多个版本
        print("未能覆盖正式文件（Word 可能正在打开 docx）：")
        print(f"  {out}")
        print(f"临时结果已保留: {tmp}")
        print("请关闭 Word 后重新运行本脚本，将直接覆盖上述正式文件。")
        sys.exit(1)
